"""
tests/test_chunker.py - Unit tests for text chunking and extraction

Tests PDF, DOCX, TXT extraction and text chunking algorithms.
"""
import pytest
import io
from unittest.mock import Mock, patch
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


class TestExtractors:
    """Test file format extractors."""

    def test_extract_text_from_txt(self):
        """Test TXT file extraction."""
        from chunker import extract_text_from_txt
        
        text = "Hello World\nThis is a test"
        data = text.encode("utf-8")
        
        result = extract_text_from_txt(data)
        assert result == text


    def test_extract_text_from_txt_with_invalid_utf8(self):
        """Test TXT extraction with invalid UTF-8 (fallback to replace)."""
        from chunker import extract_text_from_txt
        
        data = b"Hello \x80 World"
        result = extract_text_from_txt(data)
        
        # Should not raise, should replace invalid bytes
        assert isinstance(result, str)
        assert "Hello" in result
        assert "World" in result


    def test_extract_text_from_docx(self):
        """Test DOCX extraction."""
        from chunker import extract_text_from_docx
        from docx import Document
        
        # Create a minimal DOCX in memory
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        result = extract_text_from_docx(docx_bytes.read())
        assert "First paragraph" in result
        assert "Second paragraph" in result


    def test_extract_text_unsupported_format(self):
        """Test that unsupported format raises ValueError."""
        from chunker import extract_text
        
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text("document.xyz", b"some data")


class TestChunking:
    """Test text chunking algorithm."""

    def test_chunk_text_basic(self):
        """Test basic text chunking."""
        from chunker import chunk_text
        
        text = " ".join(["word"] * 100)  # 100 words
        chunks = chunk_text(text, chunk_size=10, overlap=0)
        
        assert len(chunks) == 10
        assert all(isinstance(c, str) for c in chunks)


    def test_chunk_text_with_overlap(self):
        """Test chunking with overlap."""
        from chunker import chunk_text
        
        text = " ".join([f"word{i}" for i in range(100)])
        chunks = chunk_text(text, chunk_size=20, overlap=5)
        
        # With overlap, first word of chunk N+1 should appear in chunk N
        assert "word20" in chunks[0]
        assert "word15" in chunks[1]  # 20 - 5 overlap = starts at 15


    def test_chunk_text_empty(self):
        """Test chunking empty text."""
        from chunker import chunk_text
        
        result = chunk_text("")
        assert result == []


    def test_chunk_text_whitespace_normalization(self):
        """Test that multiple spaces/newlines are normalized."""
        from chunker import chunk_text
        
        text = "word1  \n\n  word2    word3"
        chunks = chunk_text(text, chunk_size=100, overlap=0)
        
        # Should normalize to single spaces
        assert chunks[0] == "word1 word2 word3"


    def test_chunk_text_max_bytes_truncation(self):
        """Test that chunks are truncated at 3900 bytes."""
        from chunker import chunk_text
        
        # Create text that would exceed 3900 bytes
        long_word = "x" * 1000  # 1000 char word
        text = " ".join([long_word] * 10)  # 10KB text
        
        chunks = chunk_text(text, chunk_size=100, overlap=0)
        
        # All chunks should be <= 3900 bytes when encoded
        for chunk in chunks:
            assert len(chunk.encode("utf-8")) <= 3900


# TODO: Add test for extract_text_from_pdf when pymupdf setup is verified
